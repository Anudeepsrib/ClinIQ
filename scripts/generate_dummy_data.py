import os
import pandas as pd
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook

DATA_DIR = "./data/docs"
os.makedirs(DATA_DIR, exist_ok=True)

def generate_pdf():
    filepath = os.path.join(DATA_DIR, "policy_mri_authorization.pdf")
    c = canvas.Canvas(filepath)
    c.drawString(100, 800, "Enterprise Healthcare Policy: MRI Authorization")
    c.drawString(100, 780, "Policy ID: RAD-2024-001")
    c.drawString(100, 760, "Effective Date: 2024-01-01")
    
    text_lines = [
        "1. Purpose",
        "This policy defines the prior authorization requirements for Magnetic Resonance Imaging (MRI).",
        "",
        "2. Scope",
        "Applies to all outpatient MRI services for Commercial and Medicare Advantage plans.",
        "",
        "3. Medical Necessity Guidelines",
        "Prior authorization is REQUIRED for all non-emergent MRI scans.",
        "Approval is granted if ONE of the following is met:",
        " - X-ray or Ultrasound completed within the last 6 weeks showing inconclusive results.",
        " - Patient has a history of malignancy with suspicion of metastasis.",
        " - Sudden onset of neurological deficit (e.g., foot drop, cauda equina syndrome).",
        "",
        "4. Exclusions",
        " - Emergency Room visits do not require prior auth.",
        " - Inpatient stays are exempt.",
    ]
    
    y = 730
    for line in text_lines:
        c.drawString(100, y, line)
        y -= 20
        
    c.save()
    print(f"Generated PDF: {filepath}")

def generate_docx():
    filepath = os.path.join(DATA_DIR, "procedure_billing_sop.docx")
    doc = Document()
    doc.add_heading('Standard Operating Procedure: Billing & Coding', 0)
    
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        'This document outlines the coding standards for outpatient procedures.'
    )
    
    doc.add_heading('2. Modifiers', level=1)
    doc.add_paragraph('Use Modifier 25 for significant, separately identifiable evaluation and management service by the same physician on the same day of the procedure.')
    doc.add_paragraph('Use Modifier 59 for distinct procedural service.')
    
    doc.add_heading('3. Denial Management', level=1)
    doc.add_paragraph(
        'If a claim is denied for "Medical Necessity", review the attached clinical notes. '
        'Ensure that the ICD-10 code supports the CPT code billed.'
    )
    
    doc.save(filepath)
    print(f"Generated DOCX: {filepath}")

def generate_excel():
    filepath = os.path.join(DATA_DIR, "coverage_table_2024.xlsx")
    
    data = [
        {"CPT_Code": "70551", "Description": "MRI Brain w/o contrast", "Prior_Auth": "Yes", "Copay_Commercial": "$50", "Copay_Medicare": "$20", "Limit": "1 per 6 months"},
        {"CPT_Code": "70553", "Description": "MRI Brain w/ and w/o contrast", "Prior_Auth": "Yes", "Copay_Commercial": "$75", "Copay_Medicare": "$30", "Limit": "1 per year"},
        {"CPT_Code": "73721", "Description": "MRI Knee w/o contrast", "Prior_Auth": "Yes", "Copay_Commercial": "$40", "Copay_Medicare": "$15", "Limit": "None"},
        {"CPT_Code": "99213", "Description": "Office Visit Level 3", "Prior_Auth": "No", "Copay_Commercial": "$20", "Copay_Medicare": "$10", "Limit": "None"},
        {"CPT_Code": "99214", "Description": "Office Visit Level 4", "Prior_Auth": "No", "Copay_Commercial": "$25", "Copay_Medicare": "$15", "Limit": "None"},
    ]
    
    df = pd.DataFrame(data)
    df.to_excel(filepath, index=False)
    print(f"Generated Excel: {filepath}")

if __name__ == "__main__":
    generate_pdf()
    generate_docx()
    generate_excel()
